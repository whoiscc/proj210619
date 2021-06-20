from lxml import etree
from pathlib import Path
from re import finditer, escape, compile, match
from pyinflect import getAllInflections
from docx import Document


with open("words.txt") as word_file:
    word_list = [
        word.strip() for word in word_file if not word.startswith("#") and word.strip()
    ]
print(len(word_list))

re_dict = {}
for word in word_list:
    word_re = f"({escape(word)})|({escape(word.capitalize())})"
    word_ext_set = {
        word_ext
        for word_ext_list in getAllInflections(word).values()
        for word_ext in word_ext_list
    }
    for word_ext in word_ext_set:
        word_re += f"|({escape(word_ext)})|({escape(word_ext.capitalize())})"
    word_re = r"(?<![a-zA-Z])(" + word_re + r")(?![a-zA-Z])"
    re_dict[word] = compile(word_re)


def find_all(word, text):
    return list(finditer(re_dict[word], text))


i = 0
data = []
for article_path in Path("articles").iterdir():
    print(article_path)
    with open(article_path) as article_file:
        html = etree.HTML(article_file.read())
        # print(etree.tostring(html))
        # title = match(r".*Passage 0[123] (.*)", article_path.stem)[1]
        title = article_path.stem
        # print(title)
        para_list = [
            p.text.strip() for p in html.cssselect("p") if p.text and p.text.strip()
        ]

        article_data = {}
        article_data["title"] = title
        article_data["para"] = []
        appeared_set = set()
        for text in para_list:  # title?
            text_marked = []
            for word in word_list:
                if find_list := find_all(word, text):
                    appeared_set.add(word)
                    text_marked += find_list
            text_marked = sorted(text_marked, key=lambda m: m.start(0))
            # print(text_marked)
            article_data["para"].append({"text": text, "marked": text_marked})
        print(len(appeared_set))
        article_data["appeared"] = appeared_set
        data.append(article_data)

data = sorted(data, key=lambda article: len(article["appeared"]), reverse=True)
top10 = [data[0]]
del data[0]
for i in range(9):
    print(f"round {i + 1}")
    word_list = list(set(word_list) - top10[-1]["appeared"])
    for article in data:
        appeared_set = set()
        for para in article["para"]:
            for word in word_list:
                if find_all(word, para["text"]):
                    appeared_set.add(word)
        print(article["title"])
        print(len(appeared_set))
        article["appeared"] = appeared_set
    data = sorted(data, key=lambda article: len(article["appeared"]), reverse=True)
    top10.append(data[0])
    del data[0]

word_list = list(set(word_list) - top10[-1]["appeared"])

document = Document()
for article in top10:
    # print(article["title"])
    # print(article["count"])
    document.add_heading(article["title"])
    for para in article["para"]:
        p = document.add_paragraph()
        pos = 0
        for mark in para["marked"]:
            p.add_run(para["text"][pos : mark.start(0)])
            p.add_run(para["text"][max(mark.start(0), pos) : mark.end(0)]).bold = True
            pos = mark.end(0)
        p.add_run(para["text"][pos:])
    document.add_page_break()

document.add_paragraph(
    f'covered by top 10: {sum(len(article["appeared"]) for article in top10)}'
)

document.save("output.docx")

with open("uncovered.txt", "w") as uncovered_file:
    uncovered_file.write("\n".join(sorted(word_list)))
